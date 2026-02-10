# utils/export_tools.py
"""
Export Tools for Chat History and Analysis Results
Supports PDF, Word, and CSV export formats
"""

import pandas as pd
from datetime import datetime
from typing import List, Dict
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io


class ChatExporter:
    """Export chat history and analysis results"""
    
    def __init__(self):
        self.export_metadata = {
            'export_time': datetime.now(),
            'tool_name': 'AI Research Literature Assistant'
        }
    
    def export_to_pdf(
        self,
        chat_history: List[Dict],
        filename: str = "chat_export.pdf",
        include_metadata: bool = True
    ) -> str:
        """
        Export chat history to PDF
        
        Args:
            chat_history: List of chat exchanges
            filename: Output filename
            include_metadata: Whether to include document metadata
            
        Returns:
            Path to generated PDF file
        """
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, txt="Research Chat Export", ln=True, align='C')
        pdf.ln(5)
        
        # Export metadata
        pdf.set_font("Arial", 'I', 10)
        pdf.cell(0, 5, txt=f"Exported: {self.export_metadata['export_time'].strftime('%Y-%m-%d %H:%M')}", ln=True)
        pdf.ln(10)
        
        # Chat history
        for i, entry in enumerate(chat_history, 1):
            # Question
            pdf.set_font("Arial", 'B', 12)
            question_text = f"Q{i}: {entry.get('question', entry.get('query', 'N/A'))}"
            pdf.multi_cell(0, 8, txt=self._clean_text(question_text))
            pdf.ln(2)
            
            # Answer
            pdf.set_font("Arial", '', 11)
            answer_text = entry.get('answer', entry.get('response', 'N/A'))
            pdf.multi_cell(0, 7, txt=self._clean_text(answer_text))
            pdf.ln(2)
            
            # Sources if available
            if 'sources' in entry and entry['sources']:
                pdf.set_font("Arial", 'I', 9)
                sources_text = f"Sources: {', '.join(str(s) for s in entry['sources'][:3])}"
                pdf.multi_cell(0, 5, txt=self._clean_text(sources_text))
            
            # Confidence if available
            if 'confidence' in entry:
                pdf.set_font("Arial", 'I', 9)
                pdf.cell(0, 5, txt=f"Confidence: {entry['confidence']:.1f}%", ln=True)
            
            pdf.ln(5)
            
            # Add page break if needed
            if pdf.get_y() > 250:
                pdf.add_page()
        
        # Save PDF
        pdf.output(filename)
        return filename
    
    def export_to_word(
        self,
        chat_history: List[Dict],
        filename: str = "chat_export.docx",
        include_metadata: bool = True
    ) -> str:
        """
        Export chat history to Word document
        
        Args:
            chat_history: List of chat exchanges
            filename: Output filename
            include_metadata: Whether to include document metadata
            
        Returns:
            Path to generated Word file
        """
        doc = Document()
        
        # Title
        title = doc.add_heading('Research Chat Export', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        if include_metadata:
            metadata_para = doc.add_paragraph()
            metadata_para.add_run(
                f"Exported: {self.export_metadata['export_time'].strftime('%Y-%m-%d %H:%M')}"
            ).italic = True
            metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        
        # Chat history
        for i, entry in enumerate(chat_history, 1):
            # Question
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"Q{i}: {entry.get('question', entry.get('query', 'N/A'))}")
            q_run.bold = True
            q_run.font.size = Pt(12)
            q_run.font.color.rgb = RGBColor(31, 119, 180)
            
            # Answer
            a_para = doc.add_paragraph(entry.get('answer', entry.get('response', 'N/A')))
            a_para.paragraph_format.left_indent = Inches(0.25)
            
            # Sources
            if 'sources' in entry and entry['sources']:
                sources_para = doc.add_paragraph()
                sources_run = sources_para.add_run(
                    f"Sources: {', '.join(str(s) for s in entry['sources'][:3])}"
                )
                sources_run.italic = True
                sources_run.font.size = Pt(9)
                sources_para.paragraph_format.left_indent = Inches(0.25)
            
            # Confidence
            if 'confidence' in entry:
                conf_para = doc.add_paragraph()
                conf_run = conf_para.add_run(f"Confidence: {entry['confidence']:.1f}%")
                conf_run.italic = True
                conf_run.font.size = Pt(9)
                conf_para.paragraph_format.left_indent = Inches(0.25)
            
            doc.add_paragraph()
        
        # Save document
        doc.save(filename)
        return filename
    
    def export_to_csv(
        self,
        chat_history: List[Dict],
        filename: str = "chat_export.csv"
    ) -> str:
        """
        Export chat history to CSV
        
        Args:
            chat_history: List of chat exchanges
            filename: Output filename
            
        Returns:
            Path to generated CSV file
        """
        # Prepare data
        data = []
        for i, entry in enumerate(chat_history, 1):
            row = {
                'ID': i,
                'Question': entry.get('question', entry.get('query', 'N/A')),
                'Answer': entry.get('answer', entry.get('response', 'N/A')),
                'Confidence': entry.get('confidence', ''),
                'Sources': ', '.join(str(s) for s in entry.get('sources', [])),
                'Timestamp': entry.get('timestamp', self.export_metadata['export_time'])
            }
            data.append(row)
        
        # Create DataFrame and save
        df = pd.DataFrame(data)
        df.to_csv(filename, index=False, encoding='utf-8')
        return filename
    
    def export_comparison_table(
        self,
        comparison_df: pd.DataFrame,
        filename: str = "comparison.xlsx"
    ) -> str:
        """
        Export comparison table to Excel
        
        Args:
            comparison_df: DataFrame with comparison data
            filename: Output filename
            
        Returns:
            Path to generated Excel file
        """
        with pd.ExcelWriter(filename, engine='openpyxl') as writer:
            comparison_df.to_excel(writer, sheet_name='Comparison', index=False)
            
            # Auto-adjust column widths
            worksheet = writer.sheets['Comparison']
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        return filename
    
    def _clean_text(self, text: str) -> str:
        """Clean text for PDF export (remove problematic characters)"""
        if not text:
            return ""
        # Replace problematic characters
        replacements = {
            '\u2018': "'",  # Left single quote
            '\u2019': "'",  # Right single quote
            '\u201c': '"',  # Left double quote
            '\u201d': '"',  # Right double quote
            '\u2013': '-',  # En dash
            '\u2014': '-',  # Em dash
            '\u2022': '*',  # Bullet
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        # Remove non-ASCII characters
        text = text.encode('ascii', 'ignore').decode('ascii')
        return text


class SummarizationEngine:
    """Generate smart summaries at different levels"""
    
    SUMMARY_PROMPTS = {
        'executive': """Provide a concise 3-4 sentence executive summary covering:
        - Main contribution
        - Key methodology
        - Primary results
        
        Be specific and focus only on the most important information.""",
        
        'detailed': """Provide a comprehensive 1-page summary including:
        - Background and motivation
        - Methodology overview
        - Key results and findings
        - Main conclusions and limitations
        
        Use clear paragraphs and maintain academic tone.""",
        
        'sections': """Extract and summarize each major section:
        - Abstract (if present)
        - Introduction/Background
        - Methodology/Approach
        - Results/Findings
        - Conclusion/Discussion
        
        Provide 2-3 sentences for each section."""
    }
    
    @staticmethod
    def generate_summary(
        full_text: str,
        summary_type: str,
        generator_func,
        max_length: int = None
    ) -> str:
        """
        Generate summary of specified type
        
        Args:
            full_text: Full document text or chunks
            summary_type: Type of summary ('executive', 'detailed', 'sections')
            generator_func: LLM generation function
            max_length: Optional max length
            
        Returns:
            Generated summary
        """
        if summary_type not in SummarizationEngine.SUMMARY_PROMPTS:
            summary_type = 'executive'
        
        prompt = SummarizationEngine.SUMMARY_PROMPTS[summary_type]
        full_prompt = f"{prompt}\n\nDocument content:\n{full_text[:4000]}"  # Limit context
        
        try:
            summary = generator_func(full_prompt)
            
            if max_length and len(summary) > max_length:
                summary = summary[:max_length] + "..."
            
            return summary
        except Exception as e:
            return f"Error generating summary: {str(e)}"
    
    @staticmethod
    def extract_key_points(
        full_text: str,
        generator_func,
        num_points: int = 5
    ) -> List[str]:
        """
        Extract key bullet points from text
        
        Args:
            full_text: Full document text
            generator_func: LLM generation function
            num_points: Number of key points to extract
            
        Returns:
            List of key points
        """
        prompt = f"""Extract the {num_points} most important key points from this research paper.
        Format as a numbered list.
        
        Content:
        {full_text[:4000]}"""
        
        try:
            response = generator_func(prompt)
            # Parse into list
            lines = response.strip().split('\n')
            points = [line.strip() for line in lines if line.strip() and (line[0].isdigit() or line.startswith('-'))]
            return points[:num_points]
        except Exception as e:
            return [f"Error extracting key points: {str(e)}"]
